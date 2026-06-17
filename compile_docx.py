import os
import re
import sys

# Instructions to user in case of import error
try:
    import docx
    # Monkeypatch docx to handle decimal Twips measures (invalid literal for int() with base 10)
    try:
        import docx.oxml.simpletypes as st
        # Loop through all classes in simpletypes and wrap convert_from_xml
        for name, obj in list(st.__dict__.items()):
            if isinstance(obj, type) and hasattr(obj, 'convert_from_xml'):
                orig_convert = obj.convert_from_xml
                def make_wrapper(orig_fn):
                    @classmethod
                    def wrapper(cls, str_value):
                        try:
                            return orig_fn(str_value)
                        except ValueError:
                            try:
                                # Convert float string to int string (e.g. '1700.787' -> '1700')
                                int_str = str(int(float(str_value)))
                                return orig_fn(int_str)
                            except Exception:
                                return orig_fn("0")
                    return wrapper
                setattr(obj, 'convert_from_xml', make_wrapper(orig_convert))
    except Exception as e:
        print(f"Warning monkeypatching docx: {e}")

    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("=" * 60)
    print("ERRO: A biblioteca 'python-docx' nao está instalada.")
    print("Para prosseguir com a compilação do TCC, instale-a rodando:")
    print("   pip install python-docx")
    print("=" * 60)
    sys.exit(1)

def add_formatted_text(paragraph, text):
    """
    Parses simple markdown markers (**bold**, *italic*, ***bolditalic***)
    and adds formatted runs to the paragraph.
    """
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|[^*]+)')
    for part in pattern.findall(text):
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def append_markdown_file(doc, file_path):
    print(f"Processando e inserindo: {os.path.basename(file_path)}...")
    if not os.path.exists(file_path):
        print(f"Erro: Arquivo {file_path} não encontrado.")
        return

    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    while i < len(lines):
        # We check startswith on the stripped line but we want to preserve inner content
        stripped_line = lines[i].strip()
        
        # Check for code block boundary
        if stripped_line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if not stripped_line:
            i += 1
            continue

        # If inside a code block, format as preformatted text (Courier New, single spacing, indented)
        if in_code_block:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(2)
            p_format.first_line_indent = Pt(0)
            p_format.left_indent = Pt(24) # Indent the code block area
            
            run = p.add_run(lines[i]) # Use the original unstripped line to preserve whitespace
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            i += 1
            continue

        # Check for headers
        if stripped_line.startswith('# '):
            p = doc.add_paragraph()
            p.style = 'Heading 1'
            run = p.add_run(stripped_line[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            i += 1
        elif stripped_line.startswith('## '):
            p = doc.add_paragraph()
            p.style = 'Heading 2'
            run = p.add_run(stripped_line[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            i += 1
        elif stripped_line.startswith('### '):
            p = doc.add_paragraph()
            p.style = 'Heading 3'
            run = p.add_run(stripped_line[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            i += 1
        elif stripped_line.startswith('#### '):
            p = doc.add_paragraph()
            p.style = 'Heading 4'
            run = p.add_run(stripped_line[5:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            i += 1
        # Check for Markdown tables
        elif stripped_line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Remove divider rows from analysis (e.g. |:---|:---:|)
            clean_table_lines = [l for l in table_lines if not re.match(r'^\|\s*:?-+:?\s*\|', l)]
            
            rows_data = []
            for tl in clean_table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows_data.append(cells)
            
            if rows_data:
                num_cols = len(rows_data[0])
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for r_idx, row_cells in enumerate(rows_data):
                    row = table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_cells):
                        cell = row.cells[c_idx]
                        cell.text = ""  # clear default
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        add_formatted_text(p, cell_value)
                        
                        # Style run parameters
                        for run in p.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            if r_idx == 0:
                                run.bold = True  # header row
                
                # Spacing
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
        # Check for bullet items
        elif stripped_line.startswith('- '):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.5
            p_format.space_after = Pt(4)
            p_format.left_indent = Pt(18)        # Indent for bullet list
            p_format.first_line_indent = Pt(-18) # Hanging indent for bullet
            
            # Prepend bullet character
            run_bullet = p.add_run("•  ")
            run_bullet.font.name = 'Arial'
            run_bullet.font.size = Pt(12)
            
            add_formatted_text(p, stripped_line[2:])
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
            i += 1
        # Check for Figure/Table captions
        elif stripped_line.startswith('**Figura') or stripped_line.startswith('**Tabela'):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_format.line_spacing = 1.0
            p_format.space_before = Pt(8)
            p_format.space_after = Pt(2)
            p_format.first_line_indent = Pt(0)
            p_format.keep_with_next = True
            add_formatted_text(p, stripped_line)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            i += 1
        # Check for Source captions
        elif stripped_line.startswith('*Fonte:'):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_format.line_spacing = 1.0
            p_format.space_before = Pt(2)
            p_format.space_after = Pt(12)
            p_format.first_line_indent = Pt(0)
            add_formatted_text(p, stripped_line)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            i += 1
        # Check for reference and annex entries
        elif file_path.endswith('06_referencias.md') or file_path.endswith('08_anexo.md'):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_format.line_spacing = 1.0
            p_format.space_after = Pt(12)
            p_format.first_line_indent = Pt(0)
            add_formatted_text(p, stripped_line)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
            i += 1
        # Normal body paragraph (Justified, Spacing 1.5, First Line Indent 1.25cm)
        else:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_format.line_spacing = 1.5
            p_format.space_after = Pt(6)
            p_format.first_line_indent = Pt(35.4)  # 1.25 cm = 35.4 points
            add_formatted_text(p, stripped_line)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
            i += 1

def main():
    docx_input = 'TCC-Kelvin-Albuquerque-ANBT.docx'
    docx_output = 'TCC-Kelvin-Albuquerque-REVISADO.docx'
    
    revision_folder = 'docs/TCC_REVISION'
    files_to_compile = [
        '01_introducao.md',
        '02_fundamentacao.md',
        '03_metodologia.md',
        '04_resultados.md',
        '05_consideracoes.md',
        '06_referencias.md',
        '07_cronograma.md',
        '08_anexo.md'
    ]
    
    print(f"Carregando arquivo original: {docx_input}...")
    if not os.path.exists(docx_input):
        print(f"Erro: Arquivo original {docx_input} nao encontrado no diretorio root.")
        return

    doc = docx.Document(docx_input)
    
    # 1. Localizar o inicio do capitulo 1 (Introducao) para limpar tudo dali para frente
    start_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if txt == "1 introdução" or txt == "introdução" or txt == "# introdução":
            start_idx = idx
            break
            
    if start_idx == -1:
        # Tentar busca parcial por seguranca
        for idx, p in enumerate(doc.paragraphs):
            txt = p.text.strip().lower()
            if "introdução" in txt and (txt.startswith("1") or p.style.name.startswith("Heading")):
                start_idx = idx
                break

    if start_idx == -1:
        print("Erro: Nao foi possivel localizar o inicio da seção INTRODUÇÃO.")
        print("Verifique se o arquivo docx mantem o padrao de cabecalho do IFSP.")
        return

    # 2. Remover todos os elementos da introducao ate o final da body tree
    print("Limpando capítulos antigos para reinserção...")
    body = doc.element.body
    intro_p = doc.paragraphs[start_idx]
    intro_element = intro_p._element
    
    found = False
    elements_to_remove = []
    for child in body:
        if child == intro_element:
            found = True
        if found:
            elements_to_remove.append(child)
            
    for elem in elements_to_remove:
        body.remove(elem)
        
    print(f"Limpeza de {len(elements_to_remove)} elementos do corpo executada com sucesso.")

    # 3. Adicionar uma quebra de pagina para garantir que a Introducao comece em nova folha
    doc.add_page_break()

    # 4. Inserir todos os arquivos MD da revisao sequencialmente
    for filename in files_to_compile:
        full_path = os.path.join(revision_folder, filename)
        append_markdown_file(doc, full_path)
        
    # 5. Salvar o arquivo revisado
    print(f"Salvando o TCC revisado como: {docx_output}...")
    doc.save(docx_output)
    print("=" * 60)
    print("SUCESSO: Seu TCC revisado foi compilado e formatado!")
    print(f"Arquivo gerado: {docx_output}")
    print("Por favor, verifique as figuras, tabelas e sumário final.")
    print("=" * 60)

if __name__ == '__main__':
    main()
